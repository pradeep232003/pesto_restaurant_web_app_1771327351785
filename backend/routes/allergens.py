"""
Allergens — the 14 legally required FSA allergens (UK / EU).

- `CATALOG` is the immutable regulatory reference (categories + the
  named sub-items the user provided). We serve it from `/api/allergens/catalog`.
- Per menu item, we store an `allergens` dict on the `menu_items`
  document itself:
     `allergens = { category_id: [sub_item_id, ...] }`
  This keeps queries simple (one collection) and the public menu
  endpoint automatically returns the allergen data for allergen
  labelling on the customer-facing menu.
- The matrix endpoint returns every item for a location together with
  its current allergen selections so the frontend can render the grid
  in one round-trip.
"""
from typing import Dict, List, Optional
import logging

from fastapi import APIRouter, Depends, HTTPException, Query
from pydantic import BaseModel, Field

from db import db, menu_items_collection
from auth import get_admin_user, get_staff_or_above

_log = logging.getLogger("allergens")

router = APIRouter(prefix="/api/allergens", tags=["allergens"])


# ------------------------ Regulatory catalog ------------------------
# Ordered as the FSA lists them. `id` is the stable machine key,
# `label` is user-facing, `items` are the named sub-items supplied by
# the operator (Jolly's Kafe) so managers can be precise about which
# sub-item is present in a dish (e.g. "wheat" vs "spelt" for gluten).
CATALOG: List[Dict] = [
    {"id": "gluten", "label": "Cereals containing gluten", "items": [
        "wheat", "rye", "barley", "oats", "spelt", "kamut",
        "durum", "emmer", "einkorn", "triticale",
    ]},
    {"id": "crustaceans", "label": "Crustaceans", "items": [
        "prawns", "shrimp", "crab", "lobster", "crayfish",
        "langoustine", "scampi", "krill",
    ]},
    {"id": "eggs", "label": "Eggs", "items": [
        "hen_eggs", "duck_eggs", "quail_eggs", "goose_eggs",
        "turkey_eggs", "egg_white", "egg_yolk",
    ]},
    {"id": "fish", "label": "Fish", "items": [
        "salmon", "tuna", "cod", "haddock", "mackerel",
        "sardines", "trout", "anchovies", "bass", "halibut",
    ]},
    {"id": "peanuts", "label": "Peanuts", "items": [
        "peanuts", "groundnuts", "monkey_nuts", "peanut_butter",
        "peanut_flour", "peanut_oil_unrefined",
    ]},
    {"id": "soybeans", "label": "Soybeans", "items": [
        "soybeans", "edamame", "tofu", "tempeh", "soy_milk",
        "soy_flour", "soy_protein", "soy_sauce", "miso", "natto",
    ]},
    {"id": "milk", "label": "Milk", "items": [
        "cow_milk", "goat_milk", "sheep_milk", "cream", "butter",
        "cheese", "yoghurt", "whey", "casein", "buttermilk",
    ]},
    {"id": "tree_nuts", "label": "Nuts (Tree nuts)", "items": [
        "almonds", "hazelnuts", "walnuts", "cashews", "pecans",
        "brazil_nuts", "pistachios", "macadamia_nuts", "queensland_nuts",
    ]},
    {"id": "celery", "label": "Celery", "items": [
        "celery_stalks", "celery_leaves", "celeriac",
        "celery_salt", "celery_seeds",
    ]},
    {"id": "mustard", "label": "Mustard", "items": [
        "mustard_seeds_yellow", "mustard_seeds_brown", "mustard_seeds_black",
        "mustard_powder", "dijon_mustard", "wholegrain_mustard", "english_mustard",
    ]},
    {"id": "sesame", "label": "Sesame seeds", "items": [
        "sesame_seeds", "tahini", "sesame_oil_unrefined",
        "sesame_paste", "gomasio",
    ]},
    {"id": "sulphites", "label": "Sulphur dioxide & sulphites (>10 mg/kg)", "items": [
        "e220_sulphur_dioxide", "e221_sodium_sulphite", "e222_sodium_bisulphite",
        "e223_sodium_metabisulphite", "e224_potassium_metabisulphite",
        "e226_calcium_sulphite", "e227_calcium_bisulphite",
        "e228_potassium_bisulphite",
    ]},
    {"id": "lupin", "label": "Lupin", "items": [
        "lupin_beans", "lupin_flour", "lupin_seeds", "lupin_protein",
    ]},
    {"id": "molluscs", "label": "Molluscs", "items": [
        "mussels", "oysters", "clams", "scallops", "cockles",
        "squid", "octopus", "cuttlefish", "snails", "whelks", "abalone",
    ]},
]

_VALID_IDS = {c["id"]: {s for s in c["items"]} for c in CATALOG}
_VALID_CAT_IDS = set(_VALID_IDS.keys())


def _sanitise(allergens: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Silently drop unknown category / sub-item ids so a stale UI can't
    poison the DB. Returns a new dict with only valid values."""
    out: Dict[str, List[str]] = {}
    for cat_id, subs in (allergens or {}).items():
        if cat_id not in _VALID_IDS:
            continue
        valid_subs = [s for s in (subs or []) if s in _VALID_IDS[cat_id]]
        # Only persist categories that actually apply.
        if valid_subs:
            out[cat_id] = valid_subs
    return out


def _sanitise_may_contain(vals: List[str]) -> List[str]:
    """Deduplicate + drop unknown ids from the `may contain` list.
    Preserves insertion order so admin-set order shows on the print."""
    seen: set = set()
    out: List[str] = []
    for v in (vals or []):
        if v in _VALID_CAT_IDS and v not in seen:
            seen.add(v)
            out.append(v)
    return out


# ------------------------ Endpoints ------------------------
@router.get("/catalog")
async def get_catalog(user: dict = Depends(get_staff_or_above)):
    """Static 14-allergen reference — safe to cache aggressively on the
    client."""
    return {"catalog": CATALOG}


class AllergenPayload(BaseModel):
    allergens: Dict[str, List[str]] = Field(default_factory=dict)
    # Optional "may contain" — cross-contamination advisory, list of
    # top-level allergen ids (subset of _VALID_CAT_IDS). Shown on the
    # matrix + printed on the Word export as e.g. "May contain: Milk, Eggs".
    may_contain: Optional[List[str]] = None


@router.put("/matrix/{item_id}")
async def set_item_allergens(
    item_id: str,
    payload: AllergenPayload,
    user: dict = Depends(get_admin_user),
):
    existing = menu_items_collection.find_one({"id": item_id})
    if not existing:
        raise HTTPException(404, "Menu item not found")
    clean = _sanitise(payload.allergens)
    update: dict = {"allergens": clean}
    # `may_contain=None` means "don't touch" — the client didn't send it.
    # `may_contain=[]` explicitly clears the list. Sanitise either way.
    if payload.may_contain is not None:
        update["may_contain"] = _sanitise_may_contain(payload.may_contain)
    menu_items_collection.update_one(
        {"id": item_id},
        {"$set": update},
    )
    _log.info(
        "allergens: item=%s (%s) categories=%d subitems=%d may_contain=%d by=%s",
        item_id, existing.get("name", ""),
        len(clean), sum(len(v) for v in clean.values()),
        len(update.get("may_contain") or existing.get("may_contain") or []),
        user.get("email"),
    )
    return {
        "id": item_id,
        "allergens": clean,
        "may_contain": update.get("may_contain", _sanitise_may_contain(existing.get("may_contain") or [])),
    }


@router.get("/matrix")
async def get_matrix(
    location_id: str = Query(...),
    user: dict = Depends(get_staff_or_above),
):
    """Return every menu item for a location together with its current
    allergen selections — one round-trip powers the whole matrix."""
    rows = list(menu_items_collection.find(
        {"location_id": location_id},
        {"id": 1, "name": 1, "category": 1, "allergens": 1, "may_contain": 1, "is_available": 1},
    ).sort([("category", 1), ("name", 1)]).limit(2000))
    items = []
    for r in rows:
        r.pop("_id", None)
        clean = _sanitise(r.get("allergens") or {})
        may = _sanitise_may_contain(r.get("may_contain") or [])
        items.append({
            "id": r.get("id"),
            "name": r.get("name", ""),
            "category": r.get("category", ""),
            "is_available": bool(r.get("is_available", True)),
            "allergens": clean,
            "may_contain": may,
            # Convenience flag consumed by the UI to render the
            # "declared" pip on each row — saves clients from
            # counting the dict server-side on every render.
            "has_allergens": bool(clean) or bool(may),
        })
    return {"location_id": location_id, "items": items}


def _humanise(s: str) -> str:
    return (s or "").replace("_", " ").title()


@router.get("/matrix/print")
async def print_matrix(
    location_id: str = Query(...),
    user: dict = Depends(get_staff_or_above),
):
    """Generate a landscape .docx of the allergen matrix — one row per
    menu item with the declared sub-items spelt out in full so the
    document is meaningful on paper (Word/print) without needing the
    web colours + tooltips."""
    import io
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse
    from datetime import datetime, timezone

    # Pull items via the same projection used by the matrix endpoint.
    rows = list(menu_items_collection.find(
        {"location_id": location_id},
        {"id": 1, "name": 1, "category": 1, "allergens": 1, "may_contain": 1},
    ).sort([("category", 1), ("name", 1)]).limit(2000))

    doc = Document()
    # Landscape A4 with tight margins for max horizontal space.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    # Title + generation stamp.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"Allergen Matrix — {location_id}")
    run.bold = True
    run.font.size = Pt(18)

    stamp = doc.add_paragraph()
    stamp_run = stamp.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y %H:%M UTC')} · "
        f"{len(rows)} item(s) · 14 FSA allergens"
    )
    stamp_run.font.size = Pt(9)
    stamp_run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

    # Matrix table: rows = items, columns = Item | Category | 14 allergen category cells | May contain.
    header = ["Item", "Category"] + [c["label"].split(" (")[0] for c in CATALOG] + ["May contain"]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)

    # Precompute label lookup for the May-contain column.
    cat_label = {c["id"]: c["label"].split(" (")[0] for c in CATALOG}

    # One row per menu item with the actual sub-items listed in
    # each cell (or blank if the category isn't declared).
    for item in rows:
        item.pop("_id", None)
        allergens = _sanitise(item.get("allergens") or {})
        may = _sanitise_may_contain(item.get("may_contain") or [])
        row = table.add_row().cells
        row[0].text = ""
        row[0].paragraphs[0].add_run(item.get("name", "")).bold = True
        row[1].text = item.get("category", "")
        for i, cat in enumerate(CATALOG, start=2):
            subs = allergens.get(cat["id"])
            if subs:
                # All sub-items? Show ✓ instead of a long list.
                if len(subs) == len(cat["items"]):
                    row[i].text = "All"
                else:
                    row[i].text = ", ".join(_humanise(s) for s in subs)
            else:
                row[i].text = ""
        # May-contain column — rightmost.
        row[-1].text = ", ".join(cat_label.get(m, m) for m in may) if may else ""

    for tbl_row in table.rows[1:]:
        for cell in tbl_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)

    doc.add_paragraph()

    # A per-item narrative block — makes the printout usable for
    # allergy queries at the counter without deciphering the grid.
    heading = doc.add_paragraph()
    hr = heading.add_run("Detailed sub-item declarations")
    hr.bold = True
    hr.font.size = Pt(12)

    for item in rows:
        allergens = _sanitise(item.get("allergens") or {})
        may = _sanitise_may_contain(item.get("may_contain") or [])
        if not allergens and not may:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        name_run = para.add_run(f"{item.get('name', '')}")
        name_run.bold = True
        name_run.font.size = Pt(10)
        para.add_run(f"   ({item.get('category', '')})").font.size = Pt(8)
        for cat in CATALOG:
            subs = allergens.get(cat["id"])
            if not subs:
                continue
            sub_para = doc.add_paragraph(style="List Bullet")
            sub_para.paragraph_format.space_after = Pt(0)
            r1 = sub_para.add_run(f"{cat['label']}: ")
            r1.bold = True
            r1.font.size = Pt(9)
            names = ", ".join(_humanise(s) for s in subs)
            r2 = sub_para.add_run(
                "All sub-items" if len(subs) == len(cat["items"]) else names
            )
            r2.font.size = Pt(9)
        if may:
            may_para = doc.add_paragraph(style="List Bullet")
            may_para.paragraph_format.space_after = Pt(0)
            mr1 = may_para.add_run("May contain: ")
            mr1.bold = True
            mr1.font.size = Pt(9)
            mr1.font.color.rgb = RGBColor(0xA3, 0x5E, 0x00)
            mr2 = may_para.add_run(", ".join(cat_label.get(m, m) for m in may))
            mr2.font.size = Pt(9)
            mr2.font.color.rgb = RGBColor(0xA3, 0x5E, 0x00)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_loc = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in location_id)[:40]
    stamp_txt = datetime.now(timezone.utc).strftime("%Y%m%d")
    filename = f"allergen_matrix_{safe_loc}_{stamp_txt}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
