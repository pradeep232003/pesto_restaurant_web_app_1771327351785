"""
Corrective Actions Log — per-location record of failed checks and the
remedial action taken. Read-only for staff, editable by admins.

Rooted at `/api/corrective-actions`. Collection: `corrective_actions`.
Doc shape:
  {
    id, location_id, category, item, failure_description,
    corrective_action, status: 'open' | 'resolved',
    logged_by, logged_by_name, logged_at,
    resolved_by, resolved_by_name, resolved_at,
    updated_at
  }
"""
import logging
import uuid
from datetime import datetime, timezone
from typing import List, Literal, Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from pydantic import BaseModel, Field

from db import db
from auth import get_admin_user, get_staff_or_above

_log = logging.getLogger("corrective_actions")
router = APIRouter(prefix="/api/corrective-actions", tags=["corrective-actions"])

col = db["corrective_actions"]
try:
    col.create_index([("location_id", 1), ("status", 1), ("logged_at", -1)])
except Exception:
    pass


# Categories mirror the JKHive routine catalog so managers can filter
# corrective actions by the check that triggered them.
CATEGORIES = [
    "opening", "closing", "fridge_temp", "freezer_temp",
    "cooking_cooling", "reheating", "delivery", "cleaning",
    "checklist", "probe", "hygiene", "waste", "other",
]


class ActionCreate(BaseModel):
    location_id: str = Field(..., min_length=1)
    category: str = Field(..., min_length=1)
    item: str = Field("", max_length=200)
    failure_description: str = Field(..., min_length=1, max_length=1000)
    corrective_action: str = Field("", max_length=1000)
    status: Literal["open", "resolved"] = "open"


class ActionUpdate(BaseModel):
    category: Optional[str] = None
    item: Optional[str] = None
    failure_description: Optional[str] = Field(None, max_length=1000)
    corrective_action: Optional[str] = Field(None, max_length=1000)
    status: Optional[Literal["open", "resolved"]] = None


def _clean(doc: dict) -> dict:
    if not doc:
        return doc
    doc.pop("_id", None)
    return doc


@router.get("")
async def list_actions(
    location_id: str = Query(..., description="Site id — or 'all' for admin cross-site"),
    status: Literal["open", "resolved", "all"] = "all",
    limit: int = Query(500, ge=1, le=2000),
    user: dict = Depends(get_staff_or_above),
):
    """List corrective actions. Staff and above can read. Admins get an
    `'all'` shortcut for cross-site views."""
    is_admin = user.get("role") in ("admin", "super_admin")
    q: dict = {}
    if location_id and location_id != "all":
        q["location_id"] = location_id
    else:
        if not is_admin:
            raise HTTPException(400, "location_id is required for non-admin users")
    if status != "all":
        q["status"] = status
    rows = list(col.find(q).sort([("status", 1), ("logged_at", -1)]).limit(limit))
    return {"items": [_clean(r) for r in rows], "categories": CATEGORIES}


@router.post("")
async def create_action(
    payload: ActionCreate,
    user: dict = Depends(get_admin_user),
):
    """Admin-only create. Staff read but cannot mutate."""
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": str(uuid.uuid4()),
        "location_id": payload.location_id.strip(),
        "category": payload.category.strip().lower() or "other",
        "item": payload.item.strip(),
        "failure_description": payload.failure_description.strip(),
        "corrective_action": payload.corrective_action.strip(),
        "status": payload.status,
        "logged_by": user.get("email", ""),
        "logged_by_name": user.get("name") or user.get("email", ""),
        "logged_at": now,
        "resolved_by": "",
        "resolved_by_name": "",
        "resolved_at": "",
        "updated_at": now,
    }
    # If it was raised as already-resolved, stamp the resolver metadata.
    if doc["status"] == "resolved":
        doc["resolved_by"] = user.get("email", "")
        doc["resolved_by_name"] = user.get("name") or user.get("email", "")
        doc["resolved_at"] = now

    col.insert_one(doc)
    _log.info("corrective_action: created id=%s category=%s status=%s loc=%s by=%s",
              doc["id"], doc["category"], doc["status"], doc["location_id"], doc["logged_by"])
    return _clean(doc)


@router.patch("/{action_id}")
async def update_action(
    action_id: str,
    payload: ActionUpdate,
    user: dict = Depends(get_admin_user),
):
    existing = col.find_one({"id": action_id})
    if not existing:
        raise HTTPException(404, "Not found")

    updates: dict = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if payload.category is not None:
        updates["category"] = payload.category.strip().lower() or "other"
    if payload.item is not None:
        updates["item"] = payload.item.strip()
    if payload.failure_description is not None:
        updates["failure_description"] = payload.failure_description.strip()
    if payload.corrective_action is not None:
        updates["corrective_action"] = payload.corrective_action.strip()
    if payload.status is not None and payload.status != existing.get("status"):
        updates["status"] = payload.status
        if payload.status == "resolved":
            updates["resolved_at"] = datetime.now(timezone.utc).isoformat()
            updates["resolved_by"] = user.get("email", "")
            updates["resolved_by_name"] = user.get("name") or user.get("email", "")
        else:
            # Re-opened — wipe resolver so the audit trail is honest.
            updates["resolved_at"] = ""
            updates["resolved_by"] = ""
            updates["resolved_by_name"] = ""

    col.update_one({"id": action_id}, {"$set": updates})
    _log.info("corrective_action: patched id=%s fields=%s by=%s",
              action_id, list(updates.keys()), user.get("email"))
    return _clean(col.find_one({"id": action_id}))


@router.delete("/{action_id}")
async def delete_action(
    action_id: str,
    user: dict = Depends(get_admin_user),
):
    res = col.delete_one({"id": action_id})
    if res.deleted_count == 0:
        raise HTTPException(404, "Not found")
    _log.info("corrective_action: deleted id=%s by=%s", action_id, user.get("email"))
    return {"deleted": True}


@router.get("/summary")
async def summary(
    location_ids: Optional[str] = Query(None),
    user: dict = Depends(get_staff_or_above),
):
    """Cross-site open-count for dashboard badges."""
    q: dict = {"status": "open"}
    if location_ids:
        ids = [x.strip() for x in location_ids.split(",") if x.strip()]
        if ids:
            q["location_id"] = {"$in": ids}
    pipeline: List[dict] = [
        {"$match": q},
        {"$group": {"_id": "$location_id", "open": {"$sum": 1}}},
    ]
    out = {r["_id"]: r["open"] for r in col.aggregate(pipeline)}
    return {"by_location": out, "total_open": sum(out.values())}


@router.get("/print")
async def print_actions(
    location_id: str = Query(..., description="Site id — or 'all' for admin cross-site"),
    status: Literal["open", "resolved", "all"] = "all",
    days: int = Query(365, ge=1, le=3650, description="Only include entries logged in the last N days"),
    user: dict = Depends(get_staff_or_above),
):
    """Landscape .docx export of the corrective actions log. Same
    output format as the allergen matrix printer so managers know
    what to expect. Defaults to the last 12 months for EHO packs."""
    import io
    from datetime import timedelta
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    is_admin = user.get("role") in ("admin", "super_admin")
    q: dict = {}
    if location_id and location_id != "all":
        q["location_id"] = location_id
    else:
        if not is_admin:
            raise HTTPException(400, "location_id is required for non-admin users")
    if status != "all":
        q["status"] = status

    cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
    q["logged_at"] = {"$gte": cutoff}

    rows = list(col.find(q).sort([("logged_at", -1)]).limit(5000))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"Corrective Actions Log — {'All sites' if location_id == 'all' else location_id}")
    run.bold = True
    run.font.size = Pt(18)

    stamp = doc.add_paragraph()
    stamp_run = stamp.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y %H:%M UTC')} · "
        f"Last {days} day(s) · Status filter: {status} · {len(rows)} entry(ies)"
    )
    stamp_run.font.size = Pt(9)
    stamp_run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

    headers = ["Logged", "Site", "Category", "Item", "What failed", "Corrective action", "Status", "Resolved by"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    widths = [Cm(2.4), Cm(2.6), Cm(2.4), Cm(3.0), Cm(6.5), Cm(6.5), Cm(1.6), Cm(2.6)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = w
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)

    for entry in rows:
        row = table.add_row().cells
        logged_at = entry.get("logged_at", "")
        try:
            dt = datetime.fromisoformat(logged_at.replace("Z", "+00:00"))
            when = dt.strftime("%d %b %Y %H:%M")
        except Exception:
            when = logged_at[:16]
        row[0].text = when
        row[1].text = entry.get("location_id", "")
        row[2].text = _humanise_snake(entry.get("category", ""))
        row[3].text = entry.get("item", "")
        row[4].text = entry.get("failure_description", "")
        row[5].text = entry.get("corrective_action", "")
        row[6].text = (entry.get("status") or "").title()
        resolved_by = entry.get("resolved_by_name") or entry.get("resolved_by") or ""
        resolved_at = entry.get("resolved_at", "")
        try:
            if resolved_at:
                rdt = datetime.fromisoformat(resolved_at.replace("Z", "+00:00"))
                resolved_at = rdt.strftime("%d %b %Y")
        except Exception:
            resolved_at = resolved_at[:10]
        row[7].text = f"{resolved_by}\n{resolved_at}" if resolved_by else ""
        for cell, w in zip(row, widths):
            cell.width = w
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_loc = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in location_id)[:40]
    stamp_txt = datetime.now(timezone.utc).strftime("%Y%m%d")
    filename = f"corrective_actions_{safe_loc}_{stamp_txt}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _humanise_snake(s: str) -> str:
    return (s or "").replace("_", " ").title()


# ------------------------ Auto-log helper ------------------------
# Called by routine handlers when they detect a failed check. Uses a
# `source_key` (a compound like "fridge_temp:<loc>:<date>:opening:<unit>")
# so re-submissions of the same routine don't duplicate rows. Only the
# `failure_description` is updated on re-fire — if the manager has
# already added a `corrective_action` or marked the row `resolved`,
# their edits stay intact.

def auto_log_failure(
    *,
    location_id: str,
    category: str,
    item: str,
    failure_description: str,
    source_key: str,
    logged_by_email: str = "system",
    logged_by_name: str = "System (auto)",
) -> Optional[dict]:
    if not (location_id and failure_description and source_key):
        return None
    now = datetime.now(timezone.utc).isoformat()
    existing = col.find_one({"source_key": source_key})
    if existing:
        # Only refresh the failure text (data may have shifted) — never
        # touch corrective_action / status / resolved_* metadata.
        col.update_one(
            {"id": existing["id"]},
            {"$set": {
                "failure_description": failure_description[:1000],
                "item": (item or existing.get("item", ""))[:200],
                "updated_at": now,
            }},
        )
        _log.info("corrective_action: auto-refreshed id=%s key=%s", existing["id"], source_key)
        return _clean(col.find_one({"id": existing["id"]}))

    doc = {
        "id": str(uuid.uuid4()),
        "location_id": location_id,
        "category": (category or "other").strip().lower(),
        "item": (item or "")[:200],
        "failure_description": failure_description[:1000],
        "corrective_action": "",
        "status": "open",
        "source_key": source_key,
        "auto_logged": True,
        "logged_by": logged_by_email,
        "logged_by_name": logged_by_name,
        "logged_at": now,
        "resolved_by": "",
        "resolved_by_name": "",
        "resolved_at": "",
        "updated_at": now,
    }
    col.insert_one(doc)
    _log.info("corrective_action: AUTO-logged id=%s key=%s cat=%s loc=%s",
              doc["id"], source_key, doc["category"], location_id)
    return _clean(doc)


# ------------------------ Auto-resolve helper ------------------------
# Called by routine handlers when a previously-failing check now
# passes (fridge back in range, reheat now ≥ 75 °C, etc.). Closes any
# still-Open auto-logged rows for that unit/item and stamps a system
# resolve comment so the audit trail explains the closure.
#
# Match modes:
#   • Provide `source_key` to resolve exactly ONE row (event-based
#     checks such as cooking/cooling re-completion).
#   • Provide `location_id + category + item` to resolve every Open
#     row for the same unit (periodic checks such as fridge temps —
#     yesterday's opening failure closes as soon as today's opening
#     comes back in range).
#
# We only touch rows where `auto_logged=True` so admin-created rows
# stay under manual control. If a manager already typed a
# `corrective_action`, we prepend it and append the system note so
# their commentary is never lost.

def auto_resolve_failure(
    *,
    location_id: Optional[str] = None,
    category: Optional[str] = None,
    item: Optional[str] = None,
    source_key: Optional[str] = None,
    reason: str = "",
    resolver_name: str = "System (auto-resolve)",
) -> int:
    q: dict = {"status": "open", "auto_logged": True}
    if source_key:
        q["source_key"] = source_key
    else:
        if not (location_id and item):
            return 0
        q["location_id"] = location_id
        q["item"] = item[:200]
        if category:
            q["category"] = category.strip().lower()

    matches = list(col.find(q))
    if not matches:
        return 0

    now = datetime.now(timezone.utc).isoformat()
    note = f"Auto-resolved: {reason}" if reason else "Auto-resolved: reading returned to normal"
    for m in matches:
        prev = (m.get("corrective_action") or "").strip()
        final = f"{prev} | {note}" if prev else note
        col.update_one(
            {"id": m["id"]},
            {"$set": {
                "status": "resolved",
                "corrective_action": final[:1000],
                "resolved_by": "system",
                "resolved_by_name": resolver_name,
                "resolved_at": now,
                "auto_resolved": True,
                "updated_at": now,
            }},
        )
    _log.info(
        "corrective_action: AUTO-resolved %d row(s) key=%s loc=%s cat=%s item=%s",
        len(matches), source_key or "-", location_id or "-", category or "-", item or "-",
    )
    return len(matches)
